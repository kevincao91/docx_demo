import docx
import os
import shutil

ori_file_path = 'D:/PycharmProjects/docx_demo/A/'
tar_file_path = 'D:/PycharmProjects/docx_demo/B/'
muban_file_name = '典型票导入模版.docx'


def write_temp_file(temp_file_name, content_list, tar_file_name):
    docx_object = docx.Document(temp_file_name)
    #  选中表格
    table = docx_object.tables[0]
    #  清空cell
    table.rows[1].cells[1].text = ''
    #  写入需要信息
    table.rows[1].cells[1].text = content_list[0]
    for content in content_list[1:]:
        try:
            table.rows[1].cells[1].add_paragraph(content)
            # table.rows[1].cells[1].text += content
            # table.rows[1].cells[1].text += '\n'
        except Exception as e_string:
            print('写入表格失败！错误代码：', end="")
            print(e_string)
    #  查找是否有相同目标文件
    file_index = 0
    while os.path.exists(tar_file_name):
        # print(os.path.exists(tar_file_name))
        tar_file_name = tar_file_name[:-5] + str(file_index) + '.docx'
    #  写入目标文件
    try:
        docx_object.save(tar_file_name)
        print('另存为目标文件成功！')
    except Exception as e_string:
        print('另存为目标文件失败！错误代码：', end="")
        print(e_string)
    print(tar_file_name)


def read_ori_file(ori_file_name):
    docx_object = docx.Document(ori_file_name)
    tar_file_title = ''
    content_list = []
    table_index = 0
    for table in docx_object.tables:
        if table_index == 0:
            #  选中第一个表格
            #  选择需要信息
            tar_file_title = table.rows[2].cells[0].text[5:].strip()
            # print(tar_file_title)
            max_index = len(table.rows)
            #print(max_index)
            end_index = max_index - 3
            row_index = 4
            row_content = table.rows[row_index].cells[1].text.strip()
            while (row_content != 'ㄣ') & (row_index < end_index):
                # print(row_content)
                row_content = row_content.strip()
                if row_content != '':
                    content_list.append(row_content)
                row_index += 1
                row_content = table.rows[row_index].cells[1].text
        else:
            #  选中后几个表格
            #  选择需要信息
            max_index = len(table.rows)
            # print(max_index)
            end_index = max_index - 3
            row_index = 4
            row_content = table.rows[row_index].cells[1].text.strip()
            while (row_content != 'ㄣ') & (row_index < end_index):
                # print(row_content)
                row_content = row_content.strip()
                if row_content != '':
                    content_list.append(row_content)
                row_index += 1
                row_content = table.rows[row_index].cells[1].text
        #  更新table_index
        table_index += 1
    return tar_file_title, content_list


def main():
    #  生成临时文件  ===================================================================================================
    temp_file_name = tar_file_path + 'temp.docx'
    try:
        shutil.copyfile(muban_file_name, temp_file_name)
        print('模板文件拷贝到临时文件成功')
    except Exception as e_string:
        print('模板文件拷贝到临时文件失败！错误代码：', end="")
        print(e_string)
    #  遍历源文件夹所有文件  ===========================================================================================
    all_file_list = os.listdir(ori_file_path)
    total_file_num = len(all_file_list)
    index = 0
    for ori_file in all_file_list:
        index += 1
        ori_file_name = ori_file_path + ori_file
        #  读取源文件内容  获得标题 和 操作内容  =======================================================================
        tar_file_title, content_list = read_ori_file(ori_file_name)
        # print(tar_file_title)
        # print(content_list)
        tar_file_name = tar_file_path + tar_file_title + '.docx'
        #  写入临时文件内容 并另存为目标文件  ==========================================================================
        write_temp_file(temp_file_name, content_list, tar_file_name)
        string = '共%d文件，已完成第%d个文件' % (total_file_num, index)
        print(string)
    #  删除临时文件  ===================================================================================================
    os.remove(temp_file_name)
    print('临时文件清除成功！')


if __name__ == "__main__":
    main()

